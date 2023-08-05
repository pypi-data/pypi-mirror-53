import pandas as pd
from docx import Document


class QueryBuilder:

    @staticmethod
    def build_query_from_txt(file):
        f = open(file, "r")
        lines = f.readlines()
        count = lines.__len__()
        query = ''
        for l in range(0, count):
            query += '"' + lines[l].rstrip() + '"'
            if l != count - 1:
                query += ' OR '
        f.close()
        return query

    @staticmethod
    def build_query_from_word_doc(file):
        query = ''
        document = Document(file)
        paras = document.paragraphs.__len__()
        for p in range(0, paras):
            runs = document.paragraphs[p].runs.__len__()
            for r in range(0, runs):
                query += '"' + document.paragraphs[p].runs[r].text + '"'
                if r != runs - 1:
                    query += ' OR '

            if p != paras - 1:
                query += ' OR '
        return query

    @staticmethod
    def build_query_from_excel(file, sheet):
        original = pd.read_excel(file, sheet, header=None)
        query = ''
        rows = original[0].__len__()
        for row in range(0, rows):
            query += '"' + original[0][row] + '"'
            if row != rows - 1:
                query += ' OR '
        return query
