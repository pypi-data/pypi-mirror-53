import math
import pandas as pd
from docx import Document


class QueryBuilder:

    @staticmethod
    def build_query_from_txt(file):
        f = open(file, "r")
        lines = f.readlines()
        count = lines.__len__()
        query = ''
        for l in range(0, count):
            query += '"' + lines[l].rstrip() + '"'
            if l != count - 1:
                query += ' OR '
        f.close()
        return query

    @staticmethod
    def build_query_from_word_doc(file):
        query = ''
        document = Document(file)
        paras = document.paragraphs.__len__()
        for p in range(0, paras):
            runs = document.paragraphs[p].runs.__len__()
            for r in range(0, runs):
                query += '"' + document.paragraphs[p].runs[r].text + '"'
                if r != runs - 1:
                    query += ' OR '

            if p != paras - 1:
                query += ' OR '
        return query

    @staticmethod
    def build_query_from_excel(file, sheet):
        original = pd.read_excel(file, sheet, header=None)

        columns = original.columns.__len__()
        if columns == 0:
            return ''

        if columns == 1:
            return QueryBuilder.__build_simple_query_from_excel(original)

        rows = original[0].__len__()

        query = ''
        for column in range(0, columns):
            query += '('
            for row in range(0, rows):
                query += '"' + str(original[column][row]) + '"'
                if row != rows - 1:
                    if str(original[column][row+1]) == 'nan':
                        query += ')'
                        break
                    else:
                        query += ' OR '
                else:
                    query += ')'

            if column != columns - 1:
                query += ' AND '
        return query

    @staticmethod
    def __build_simple_query_from_excel(dataframe):
        query = ''
        rows = dataframe[0].__len__()
        for row in range(0, rows):
            query += '"' + str(dataframe[0][row]) + '"'
            if row != rows - 1:
                query += ' OR '
        return query

